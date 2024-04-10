import csv
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate
from fastapi import FastAPI, UploadFile, File
from typing import List
import pandas as pd
import os

template = DocxTemplate('certificate-template.docx')
output_folder = 'certifications/'

app = FastAPI()

def convert_xlsx_to_csv(input_file, output_file):
    df = pd.read_excel(input_file)
    df.to_csv(output_file, index=False)

def create_certificate(data):
    for row in data[1:]:
        if len(row) < 10:
            print(f"Skipping row {row} due to insufficient data")
            continue
        
        Sl_No, Student_1st_Name, Student_2nd_Name, Guardian_1st_Name, Guardian_2nd_Name, Course_Name, Reg_No, Academic_Session, naac_nio, photo, *extra_fields = row
        
        context = {
            'Sl_No': Sl_No,
            'Student_1st_Name': Student_1st_Name,
            'Student_2nd_Name': Student_2nd_Name,
            'Guardian_1st_Name': Guardian_1st_Name,
            'Guardian_2nd_Name': Guardian_2nd_Name,
            'Course_Name': Course_Name,
            'Reg_No': Reg_No,
            'Academic_Session': Academic_Session,
            'naac_nio': naac_nio,
            'photo': photo if photo else 'nil'  
        }
        docx_file = os.path.join(output_folder, f"{Student_1st_Name}.docx")
        template.render(context)
        template.save(docx_file)

def add_image_to_docx(docx_file, image_path):
    document = Document(docx_file)
    last_paragraph = document.paragraphs[-1]  

    run = last_paragraph.add_run()
    run.add_picture(image_path, width=Inches(2)) 

    document.save(docx_file)

@app.post('/upload')
async def handle_upload(files: List[UploadFile] = File(...)):
    for uploaded_file in files:
        if uploaded_file.filename.endswith('.xlsx'):
            contents = await uploaded_file.read()
            input_file = 'input.xlsx'
            output_file = 'output.csv'
            with open(input_file, 'wb') as f:
                f.write(contents)
            convert_xlsx_to_csv(input_file, output_file)
            df = pd.read_csv(output_file)
            csv_data = [df.columns.tolist()] + df.values.tolist()
            create_certificate(csv_data)
            return {'message': 'Certificates generated successfully'}
    
    return {'error': 'Please upload an XLSX file'}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
